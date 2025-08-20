import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
import io

# --- Dicionário de Cores e Funções Auxiliares (mantidos como no original) ---
cores_por_tema = {
    "CONHECIMENTO E INOVAÇÃO": "4400FF",
    "SAÚDE E QUALIDADE DE VIDA": "ED282C",
    "SEGURANÇA E CIDADANIA": "FFB000",
    "DESENVOLVIMENTO SUSTENTÁVEL": "87D200",
    "Gestão, Transparência e Participação": "002060"
}

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.replace("#", ""))
    tcPr.append(shd)

def set_paragraph_background(paragraph, color):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color.replace("#", ""))
    pPr.append(shd)

# --- FUNÇÃO PRINCIPAL CORRIGIDA ---
def criar_documentos_por_tema(df: pd.DataFrame):
    # --- Preparação do DataFrame (sem alterações) ---
    df.rename(columns={
        'N2 - Secretaria': 'Orgao', 'N2 - Nome': 'Iniciativa', 'N2 - Status Informado': 'Status_Informado',
        'N2 - Ação Orçamentária': 'Acao', 'N2 - Programa Orçamentário': 'Programa', 'N2 - Início Realizado': 'Inicio_Realizado',
        'N2 - Término Realizado': 'Termino_Realizado', 'N2 - Resultado': 'RGS_GGGE',
        'N2 - Localização Geográfica': 'Localizacao_Geografica', 'N1 - Nome': 'Objetivo_Estrategico'
    }, inplace=True, errors='ignore')
    
    colunas_necessarias = ['Orgao', 'Iniciativa', 'Status_Informado', 'Acao', 'Programa',
                           'Inicio_Realizado', 'Termino_Realizado', 'RGS_GGGE',
                           'Localizacao_Geografica', 'Objetivo_Estrategico']
    df2 = df[colunas_necessarias].copy()

    df2[['Inicio_Realizado', 'Termino_Realizado']] = df2[['Inicio_Realizado', 'Termino_Realizado']].apply(
        lambda x: pd.to_datetime(x, errors='coerce', dayfirst=True)
    )
    status_filter = ['EM EXECUÇÃO', 'CONCLUÍDO', 'EM LICITAÇÃO', 'LICITAÇÃO CONCLUÍDA', 'OBRA EM LICITAÇÃO']
    df2 = df2[df2['Status_Informado'].isin(status_filter)]
    status_mapping = {'EM LICITAÇÃO': 'EM EXECUÇÃO', 'LICITAÇÃO CONCLUÍDA': 'EM EXECUÇÃO', 'OBRA EM LICITAÇÃO': 'EM EXECUÇÃO'}
    df2['Status_Informado'] = df2['Status_Informado'].replace(status_mapping)

    documentos_gerados = {}
    temas_unicos = df2['Objetivo_Estrategico'].unique()

    for tema in temas_unicos:
        if pd.isna(tema):
            continue

        # A ORDENAÇÃO É A CHAVE PARA A LÓGICA FUNCIONAR CORRETAMENTE
        df_tema = df2[df2['Objetivo_Estrategico'] == tema].sort_values(
            by=['Orgao', 'Programa', 'Acao']
        )
        
        if df_tema.empty:
            continue

        doc = Document()
        section = doc.sections[0]
        section.left_margin = Inches(0)
        section.right_margin = Inches(0)
        
        cor = cores_por_tema.get(tema, "D3D3D3")

        # --- VARIÁVEIS DE CONTROLE DE ESTADO ---
        orgao_anterior = None
        programa_anterior = None
        acao_anterior = None

        for row in df_tema.itertuples(index=False):
            # --- NÍVEL 1: Checa mudança de ÓRGÃO ---
            if row.Orgao != orgao_anterior:
                if orgao_anterior is not None:
                    doc.add_page_break()
                
                # Imprime cabeçalho do Órgão
                p_orgao = doc.add_paragraph()
                p_orgao.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_orgao.paragraph_format.space_after = Pt(0) 
                run = p_orgao.add_run(str(row.Orgao).upper())
                run.font.name = 'Gilroy ExtraBold'
                run.font.size = Pt(12)
                run.bold = True
                run.font.color.rgb = RGBColor(0, 32, 96)
                set_paragraph_background(p_orgao, 'D3D3D3')
                
                # Atualiza estado e REINICIA os níveis inferiores
                orgao_anterior = row.Orgao
                programa_anterior = None
                acao_anterior = None

            # --- NÍVEL 2: Checa mudança de PROGRAMA ---
            # Esta verificação é acionada se o programa atual for diferente do anterior
            # OU se o órgão mudou (pois programa_anterior foi reiniciado para None)
            if row.Programa != programa_anterior:
                p_programa = doc.add_paragraph()
                p_programa.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_programa.paragraph_format.space_before = Pt(0)
                p_programa.paragraph_format.space_after = Pt(0) # Sem espaço para a ação vir junto

                run_programa = p_programa.add_run(str(row.Programa).upper())
                run_programa.font.name = 'Gilroy ExtraBold'
                run_programa.font.size = Pt(12)
                run_programa.bold = True
                run_programa.font.color.rgb = RGBColor(255, 255, 255)
                set_paragraph_background(p_programa, cor)

                # Atualiza estado e REINICIA o nível inferior
                programa_anterior = row.Programa
                acao_anterior = None

            # --- NÍVEL 3: Checa mudança de AÇÃO ---
            # Acionada se a ação for diferente da anterior
            # OU se o programa/órgão mudou (pois acao_anterior foi reiniciado)
            if row.Acao != acao_anterior:
                p_acao = doc.add_paragraph()
                p_acao.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_acao.paragraph_format.space_before = Pt(0)
                p_acao.paragraph_format.space_after = Pt(8) # Espaço depois da ação

                run_acao = p_acao.add_run(str(row.Acao).title())
                run_acao.font.name = 'Gilroy Light'
                run_acao.font.size = Pt(12)
                run_acao.font.color.rgb = RGBColor(255, 255, 255)
                # O fundo da ação deve ter a mesma cor do programa
                set_paragraph_background(p_acao, cor)

                # Atualiza estado
                acao_anterior = row.Acao

            # --- NÍVEL 4: Imprime os detalhes da INICIATIVA (sempre) ---
            status_imagem = 'imagens\concluido.png' if row.Status_Informado == 'CONCLUÍDO' else 'imagens\em_execucao.png'
            status_texto_label = 'Data de Entrega:' if row.Status_Informado == 'CONCLUÍDO' else 'Data de Início:'
            prazo = row.Termino_Realizado if row.Status_Informado == 'CONCLUÍDO' else row.Inicio_Realizado
            icone_localizacao_path = 'imagens\localizacao.png'
            icone_calendario_path = 'imagens\calendario.png'

            table = doc.add_table(rows=4, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False

            # (O restante do código que cria a tabela permanece exatamente o mesmo)
            cell_iniciativa = table.cell(0, 0).merge(table.cell(0, 4))
            p_iniciativa = cell_iniciativa.paragraphs[0]
            p_iniciativa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_iniciativa.add_run(str(row.Iniciativa))
            run.font.name = 'Gilroy ExtraBold'
            run.font.size = Pt(10)
            run.bold = True
            cell_iniciativa.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background(cell_iniciativa, 'D3D3D3')

            cell_status = table.cell(1, 0).merge(table.cell(1, 1))
            run_label_status = cell_status.paragraphs[0].add_run()
            run_label_status.add_picture(status_imagem, width=Inches(0.17))
            run_label_status.add_text('  Status:  ')
            run_label_status.font.name = 'Neutro Thin'
            run_label_status.font.size = Pt(9)
            run_valor_status = cell_status.paragraphs[0].add_run(str(row.Status_Informado).capitalize())
            run_valor_status.font.name = 'Neutro'
            run_valor_status.font.size = Pt(10)
            
            cell_data_merged = table.cell(1, 2).merge(table.cell(1, 4))
            run_data_label = cell_data_merged.paragraphs[0].add_run()
            run_data_label.add_picture(icone_calendario_path, width=Inches(0.17))
            run_data_label.add_text(f'  {status_texto_label} ')
            run_data_label.font.name = 'Neutro'
            run_data_label.font.size = Pt(9)
            data_texto = prazo.strftime('%d/%m/%Y') if pd.notnull(prazo) else ''
            run_data_valor = cell_data_merged.paragraphs[0].add_run(f'{data_texto}')
            run_data_valor.font.name = 'Neutro'
            run_data_valor.font.size = Pt(10)

            # --- CORREÇÃO: Linha de Localização em uma única célula com fontes diferentes ---

            # 1. Mesclar todas as células da linha para criar um único contêiner
            cell_municipios = table.cell(2, 0).merge(table.cell(2, 4))
            p_municipios = cell_municipios.paragraphs[0]
            
            # 2. Adicionar o PRIMEIRO RUN (ícone e rótulo) com a primeira fonte
            run_label = p_municipios.add_run()
            run_label.add_picture(icone_localizacao_path, width=Inches(0.17))
            run_label.add_text('  Municípios Atendidos: ') # Adicionei um espaço no final para separar
            run_label.font.name = 'Neutro Thin'
            run_label.font.size = Pt(9)

            # 3. Adicionar o SEGUNDO RUN (valor da localização) com a segunda fonte
            localizacao_texto = "" if pd.isnull(row.Localizacao_Geografica) else str(row.Localizacao_Geografica)
            run_valor = p_municipios.add_run(localizacao_texto)
            run_valor.font.name = 'Neutro'
            run_valor.font.size = Pt(10)

            # 4. (Opcional) Alinhar o conteúdo da célula verticalmente
            cell_municipios.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            cell_rgs = table.cell(3, 0).merge(table.cell(3, 4))
            run_rgs = cell_rgs.paragraphs[0].add_run(str(row.RGS_GGGE))
            run_rgs.font.name = 'Neutro'
            run_rgs.font.size = Pt(9)

        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        documentos_gerados[tema] = doc_io

    return documentos_gerados